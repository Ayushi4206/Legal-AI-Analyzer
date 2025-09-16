import os
import logging
from typing import Optional
import fitz  # PyMuPDF for PDF processing
from docx import Document  # python-docx for DOCX processing
import tempfile

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles document parsing and text extraction from various file formats
    """
    
    def __init__(self):
        """Initialize the document processor"""
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from uploaded document
        
        Args:
            file_path: Path to the uploaded file
            
        Returns:
            Extracted text content
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text_content = []
            
            # Open PDF document
            pdf_document = fitz.open(file_path)
            
            # Extract text from each page
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():  # Only add non-empty pages
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            
            pdf_document.close()
            
            extracted_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX/DOC file
        
        Args:
            file_path: Path to DOCX/DOC file
            
        Returns:
            Extracted text
        """
        try:
            # Load document
            doc = Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
    
    def get_document_metadata(self, file_path: str) -> dict:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dictionary containing metadata
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            metadata = {
                'file_size': os.path.getsize(file_path),
                'file_extension': file_extension,
                'page_count': 0
            }
            
            if file_extension == '.pdf':
                pdf_document = fitz.open(file_path)
                metadata.update({
                    'page_count': pdf_document.page_count,
                    'title': pdf_document.metadata.get('title', ''),
                    'author': pdf_document.metadata.get('author', ''),
                    'subject': pdf_document.metadata.get('subject', ''),
                    'creator': pdf_document.metadata.get('creator', '')
                })
                pdf_document.close()
                
            elif file_extension in ['.docx', '.doc']:
                doc = Document(file_path)
                metadata.update({
                    'page_count': len(doc.paragraphs),
                    'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                    'table_count': len(doc.tables)
                })
                
                # Try to get core properties
                try:
                    core_props = doc.core_properties
                    metadata.update({
                        'title': core_props.title or '',
                        'author': core_props.author or '',
                        'subject': core_props.subject or '',
                        'created': str(core_props.created) if core_props.created else '',
                        'modified': str(core_props.modified) if core_props.modified else ''
                    })
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {'error': str(e)}
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        try:
            # Remove excessive whitespace
            lines = [line.strip() for line in text.split('\n')]
            lines = [line for line in lines if line]  # Remove empty lines
            
            # Join lines with proper spacing
            cleaned_text = '\n'.join(lines)
            
            # Remove multiple consecutive spaces
            import re
            cleaned_text = re.sub(r' +', ' ', cleaned_text)
            
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Error cleaning text: {str(e)}")
            return text
    
    def split_into_sections(self, text: str) -> list:
        """
        Split document text into logical sections
        
        Args:
            text: Document text
            
        Returns:
            List of text sections
        """
        try:
            # Common legal document section headers
            section_patterns = [
                r'\b(?:SECTION|Section)\s+\d+',
                r'\b(?:ARTICLE|Article)\s+\d+',
                r'\b(?:CLAUSE|Clause)\s+\d+',
                r'\b\d+\.\s+[A-Z][A-Za-z\s]+',
                r'\b[A-Z\s]{3,}\s*$'  # All caps headers
            ]
            
            import re
            sections = []
            current_section = ""
            
            for line in text.split('\n'):
                is_header = False
                for pattern in section_patterns:
                    if re.search(pattern, line):
                        is_header = True
                        break
                
                if is_header and current_section.strip():
                    sections.append(current_section.strip())
                    current_section = line + '\n'
                else:
                    current_section += line + '\n'
            
            # Add the last section
            if current_section.strip():
                sections.append(current_section.strip())
            
            # If no sections found, split by paragraphs
            if len(sections) <= 1:
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                return paragraphs[:10]  # Limit to first 10 paragraphs
            
            return sections
            
        except Exception as e:
            logger.error(f"Error splitting into sections: {str(e)}")
            return [text]  # Return original text as single section
    
    def validate_document(self, file_path: str) -> dict:
        """
        Validate if the document can be processed
        
        Args:
            file_path: Path to document file
            
        Returns:
            Validation results
        """
        try:
            result = {
                'valid': True,
                'errors': [],
                'warnings': []
            }
            
            # Check if file exists
            if not os.path.exists(file_path):
                result['valid'] = False
                result['errors'].append('File does not exist')
                return result
            
            # Check file size (max 50MB)
            file_size = os.path.getsize(file_path)
            if file_size > 50 * 1024 * 1024:  # 50MB
                result['valid'] = False
                result['errors'].append('File size exceeds 50MB limit')
            
            # Check file extension
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension not in self.supported_formats:
                result['valid'] = False
                result['errors'].append(f'Unsupported file format: {file_extension}')
            
            # Try to extract a small sample to verify file integrity
            try:
                sample_text = self.extract_text(file_path)
                if len(sample_text.strip()) < 100:
                    result['warnings'].append('Document contains very little text content')
            except Exception as e:
                result['valid'] = False
                result['errors'].append(f'Cannot read document: {str(e)}')
            
            return result
            
        except Exception as e:
            return {
                'valid': False,
                'errors': [f'Validation error: {str(e)}'],
                'warnings': []
            }